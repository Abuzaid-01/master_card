"""
Generates the official technical submission document in Word (.docx) format
for the Mastercard Innovation Challenge @ GFF 2026.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))


def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def create_submission_document(filename="SENTRIX_AI_Mastercard_Submission.docx"):
    doc = Document()
    
    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ── Title & Header ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("SENTRIX AI")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Autonomous Multi-Modal Red Teaming & Self-Healing Defense Platform for Next-Gen Payment Security")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Mastercard Innovation Challenge @ Global Fintech Fest (GFF) 2026 | Submission Category: AI Defense Lab")
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(10)
    meta_run.font.bold = True
    meta_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("―" * 55).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Team Information Box ──
    team_table = doc.add_table(rows=3, cols=2)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_table.autofit = False
    
    headers = ["Project & Repository Links", "Submission & Team Metadata"]
    for i, h in enumerate(headers):
        cell = team_table.cell(0, i)
        cell.paragraphs[0].add_run(h).font.bold = True
        set_cell_background(cell, "F2F2F2")

    team_table.cell(1, 0).paragraphs[0].add_run("• Live Web Platform: https://sentrix-ai-seven.vercel.app/\n• FastAPI API Docs: https://master-card.onrender.com/docs\n• Public GitHub Repo: https://github.com/Abuzaid-01/master_card")
    team_table.cell(1, 1).paragraphs[0].add_run("• Team Name: Abuzaid_2005\n• Target Event: GFF 2026, Mumbai (Jio World Centre)\n• Track: Red Team / Blue Team AI Defense Lab\n• Submission Deadline: 31st August 2026")
    
    team_table.cell(2, 0).paragraphs[0].add_run("Team Members & Registered Email IDs:").font.bold = True
    team_table.cell(2, 1).paragraphs[0].add_run("1. Abuzaid — abuzaid0205@gmail.com\n2. Dipesh Kumar — dipeshkumar0853822@gmail.com")
    set_cell_background(team_table.cell(2, 0), "FAFAFA")
    set_cell_background(team_table.cell(2, 1), "FAFAFA")

    doc.add_paragraph()

    # ── Section 1: Executive Summary ──
    h1 = doc.add_heading("1. Executive Summary & Core Innovation", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Generative AI has fundamentally reshaped financial cybercrime by lowering attack costs, enabling dynamic behavioral evasion, and weaponizing conversational interfaces at scale. Traditional payment defenses operate in isolated silos—the chatbot safety team, payment authorization gateway, and anti-money laundering (AML) graph units evaluate events independently. Modern fraud syndicates exploit these blind spots via compound kill-chains: using conversational prompt injection to suppress SMS alerts, executing automated card testing micro-bursts across merchant channels, and instantly laundering funds through multi-hop mule networks within seconds."
    )
    doc.add_paragraph(
        "SENTRIX AI is an end-to-end, closed-loop Red Team/Blue Team prototype for payment networks. It combines: (1) a 36-vector fraud taxonomy informed by regulatory and real-world payment-fraud patterns; (2) synthetic tabular, text, graph, and adversarial-evasion datasets checked against explicit domain rules; and (3) specialised detection models, risk fusion, cost-aware thresholding, and active-learning retraining. The goal is straightforward: simulate coordinated fraud safely, detect it across modalities, and continuously test where the defence is weakest."
    )

    # ── Section 2: Threat Taxonomy (Pillar 1 - Identify) ──
    h2 = doc.add_heading("2. Pillar 1: Threat Identification & 36-Vector Taxonomy", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Rather than focusing on narrow fraud patterns, SENTRIX AI establishes an exhaustive 36-vector threat taxonomy categorized across 5 distinct operational pillars. Every vector is grounded in real payment protocols, cardholder behavioral data, and official regulatory advisories:"
    )

    pillar_summary = [
        ("Pillar 1: AI Red-Teaming & Social Engineering (8 Vectors)", "V01: Admin Impersonation Override, V02: API Tool Hijack, V03: AML Officer Pretext, V04: Voice Clone Vishing, V05: Safe-Account Coercion, V06: Pig-Butchering Romance, V07: B2B Invoice Forgery, V08: Agentic MCP Tool Hijack."),
        ("Pillar 2: Synthetic Data & Obfuscation (7 Vectors)", "V09: Base64/Unicode Masking, V10: Indirect Memo Injection, V11: Multi-Turn Context Poisoning, V12: Multilingual Cross-Lingual Evasion, V13: Prompt Leaking Recon, V14: Social Urgency Hijack, V15: DAN Persona Jailbreak."),
        ("Pillar 3: Multi-Rail & Digital Payment Exploits (7 Vectors)", "V16: Digital Wallet Push Provisioning Bypass, V17: Ghost Tap & Contactless NFC Relay, V18: BOPIS Mule Laundering, V19: BNPL Synthetic Stacking, V20: Refund-as-a-Service (RaaS), V21: Distributed Low-and-Slow BIN Enumeration, V22: Synthetic Merchant Bust-Out."),
        ("Pillar 4: Money Laundering & Network Topologies (7 Vectors)", "V23: Multi-Hop Linear Mule Chains, V24: Fan-Out Dispersal Sweeps, V25: Smurfing & Structuring Funnels, V26: Round-Trip Wash Cycles, V27: Instant Micro-Smurfing (UPI/FedNow), V28: Chameleon Mule Networks (90/10 noise), V29: Crypto Off-Ramp Mixers."),
        ("Pillar 5: Adversarial Evasion & Active Learning (7 Vectors)", "V30: Amount Skirting ($1.99), V31: Velocity Dilution (Poisson throttling), V32: Canvas & Residential Geo-Spoofing, V33: Decline Masking & Age Inflation, V34: Training Set Model Poisoning, V35: Adaptive RL Boundary Prober, V36: Black-Box Surrogate Boundary Gradient Probing.")
    ]

    t_pillar = doc.add_table(rows=len(pillar_summary)+1, cols=2)
    t_pillar.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_pillar.cell(0, 0).paragraphs[0].add_run("Operational Threat Pillar").font.bold = True
    t_pillar.cell(0, 1).paragraphs[0].add_run("Vectors & Real-World Threat Mechanics").font.bold = True
    set_cell_background(t_pillar.cell(0, 0), "EAEAEA")
    set_cell_background(t_pillar.cell(0, 1), "EAEAEA")

    for idx, (p_title, p_desc) in enumerate(pillar_summary):
        r_idx = idx + 1
        t_pillar.cell(r_idx, 0).paragraphs[0].add_run(p_title).font.bold = True
        t_pillar.cell(r_idx, 1).paragraphs[0].add_run(p_desc)

    doc.add_paragraph()

    # ── Section 3: Synthetic Generation & Fidelity (Pillar 2 - Generate) ──
    h3 = doc.add_heading("3. Pillar 2: Synthetic Attack Generation & Statistical Fidelity", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Each generator starts from known fraud mechanics, varies the signals an attacker would manipulate, and validates the exported records before they enter model training. This creates realistic enough stress cases without using customer data."
    )

    methodology = doc.add_paragraph()
    methodology.add_run("How vectors are produced and used. ").bold = True
    methodology.add_run(
        "The tabular generator mixes ordinary spend, hard negatives, fraud sub-types, and controlled label noise across 15 features. "
        "The text generator expands seeded prompts with template and slot variations. The graph generator builds organic transfers alongside mule rings, then derives degree, funnel, and pass-through signals. "
        "A separate evasion pass mutates fraud records to imitate threshold skirting, spoofed device signals, velocity dilution, poisoning, and boundary probing. Finally, 100 cross-vector scenarios link text, payment, and laundering stages into one campaign."
    )

    gen_stats = [
        ("Tabular Card Testing", "50,000 transactions", "15 signals covering amount, velocity, device risk, decline state, timing, merchant risk, geography, card history, and newer payment-rail indicators. Includes 12 fraud sub-types, normal activity, and hard negatives."),
        ("Prompt Injection Text", "1,500 prompts", "17 malicious categories plus legitimate customer queries. Seed prompts are expanded with deterministic templates and slot variations, including multilingual, encoded, multi-turn, and tool-use attacks."),
        ("Money Mule Network Graph", "7,478 transfers", "1,000 accounts and 100 mule rings across seven laundering patterns: chains, fan-out, smurfing, wash cycles, instant micro-smurfing, chameleon traffic, and crypto off-ramping."),
        ("Adversarial Evasion", "1,211 mutated fraud records", "Seven perturbation strategies modify a copy of the tabular data: amount structuring, timing jitter, device spoofing, velocity dilution, poisoning triggers, and two forms of boundary probing."),
        ("Cross-Vector Scenarios", "100 Unique Campaigns", "Procedural multi-stage campaigns correlating chatbot text injection, gateway card testing, and graph laundering flows with 100% unique entity IDs.")
    ]

    t_gen = doc.add_table(rows=len(gen_stats)+1, cols=3)
    t_gen.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col_name in enumerate(["Modality", "Volume", "How It Is Generated"]):
        c = t_gen.cell(0, i)
        c.paragraphs[0].add_run(col_name).font.bold = True
        set_cell_background(c, "EAEAEA")

    for idx, (m, v, f) in enumerate(gen_stats):
        r = idx + 1
        t_gen.cell(r, 0).paragraphs[0].add_run(m).font.bold = True
        t_gen.cell(r, 1).paragraphs[0].add_run(v)
        t_gen.cell(r, 2).paragraphs[0].add_run(f)

    doc.add_paragraph(
        "\nValidation note: the project also runs Train-on-Synthetic, Test-on-Real (TSTR) and distributional checks against 20,000 IEEE-CIS benchmark records. These results are an early utility and distributional comparison—not a claim that synthetic data is equivalent to production traffic."
    )

    # ── Section 4: Multi-Model Defense Heads (Pillar 3 - Defend) ──
    h4 = doc.add_heading("4. Pillar 3: Multi-Model Defense Heads & Efficacy", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The defence suite uses three specialist heads in parallel. Each returns a risk probability; the fusion layer combines the signals and selects an operational recommendation."
    )

    def_metrics = [
        ("Tabular Fraud Detector", "ONNX-exported XGBoost + Isolation Forest", "AUC-PR: 0.8031\nF1 Score: 0.8147\nFPR: 0.0035 (0.35%)", "0.006 ms", "Meets the project’s sub-50 ms inference target on the local benchmark"),
        ("NLP Injection Detector", "ONNX MiniLM embeddings + calibrated classifier and similarity ensemble", "Semantic AUC-PR: 1.0000\nParaphrased lift: +9.19% vs TF-IDF\nThreshold: 0.2927", "Not separately reported", "Captures meaning-level variation that keyword-only features can miss"),
        ("Mule Graph Classifier", "HistGradientBoosting on NetworkX topology features", "AUC-PR: 0.9630\nF1 Score: 0.9025\nThreshold tuned on validation", "Not separately reported", "Scores rapid pass-through, high-degree, and funnel-like behaviour"),
        ("Financial Cost Optimizer", "Asymmetric Loss Sweeper ($1.2*Amt FN vs $15.00 FP)", "Optimal Tau: 0.36\nMin Financial Loss: $46,218.98\nSavings vs Default: $3,071.97", "Pre-computed", "Reduces chargeback losses while maintaining frictionless commerce"),
        ("Cross-Vector Fusion Engine", "Joint-risk formula + action policy", "1 - Product(1-Pi) + synergy\nCritical: >=0.80\nHigh: 0.50–0.79", "Scenario execution", "Returns freeze, step-up, or monitor recommendations for payment-rail integration")
    ]

    t_def = doc.add_table(rows=len(def_metrics)+1, cols=5)
    t_def.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col_name in enumerate(["Defense Head", "Model Architecture", "Benchmark Accuracy", "Latency", "Enterprise Impact"]):
        c = t_def.cell(0, i)
        c.paragraphs[0].add_run(col_name).font.bold = True
        set_cell_background(c, "EAEAEA")

    for idx, (dh, ma, ba, lat, imp) in enumerate(def_metrics):
        r = idx + 1
        t_def.cell(r, 0).paragraphs[0].add_run(dh).font.bold = True
        t_def.cell(r, 1).paragraphs[0].add_run(ma)
        t_def.cell(r, 2).paragraphs[0].add_run(ba)
        t_def.cell(r, 3).paragraphs[0].add_run(lat)
        t_def.cell(r, 4).paragraphs[0].add_run(imp)

    doc.add_paragraph(
        "\nExplainability: global and per-transaction TreeSHAP attribution is available across all 15 tabular features (the top three are failed_attempts_24h, nfc_tap_latency_ms, and bnpl_bureau_inquiries). This provides traceable model evidence to support a future PCI-DSS and model-governance review; it is not a compliance certification."
    )

    # ── Section 5: Closed-Loop Active Learning Engine ──
    h5 = doc.add_heading("5. Pillar 4: Closed-Loop Adversarial Active Learning", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The closed loop makes the evaluation practical: it keeps an unseen partition, probes the first-round models, retrains on discovered misses, and then checks both adversarial and baseline performance:"
    )

    loop_steps = [
        "1. Holdout Partitioning: Step 4 holdout data is split 50/50 into a mining partition and an unseen final-evaluation partition.",
        "2. Model-Aware Probing: targeted tabular, graph, and text mutations probe the Round 1 decision boundaries to expose blind spots.",
        "3. Failure Mining and Retraining: evaded samples are automatically labelled and added to the training pool (+4,683 tabular and +711 graph records).",
        "4. Evaluation: on unseen adversarial graph data, the mule catch rate improved from 39.8% to 87.5% (+42 catches) and AUC-PR rose from 0.7114 to 0.9404. Baseline graph AUC-PR changed by only -0.0019, so the project’s forgetting check did not flag material regression."
    ]

    for step in loop_steps:
        doc.add_paragraph(step)

    # ── Section 6: Conclusion & Deployment ──
    h6 = doc.add_heading("6. Conclusion & Live Deployment Links", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "SENTRIX AI demonstrates a unified prototype that generates, stress-tests, and improves fraud defences across text, payment, and money-movement signals. The source, live web demonstration, and API documentation are available here:"
    )

    doc.add_paragraph("• Interactive Web Dashboard: https://sentrix-ai-seven.vercel.app/")
    doc.add_paragraph("• Live FastAPI Backend & API Docs: https://master-card.onrender.com/docs")
    doc.add_paragraph("• GitHub Public Source Code: https://github.com/Abuzaid-01/master_card")

    # Save to disk
    out_path = os.path.join(PROJECT_ROOT, filename)
    doc.save(out_path)
    print(f"[Success] Generated Word Submission Document: {out_path}")
    return out_path


if __name__ == "__main__":
    create_submission_document("Abuzaid_2005.docx")
    create_submission_document("Mastercard_Innovation_Challenge_Submission_Report.docx")
    create_submission_document("TeamName.docx")
